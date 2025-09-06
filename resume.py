pip install python-docx reportlab

from docx import Document
from docx.shared import Pt
import docx2pdf

# Naya Word document banao
doc = Document()

# Title
doc.add_heading("Jairaj Choudhary", 0)
doc.add_paragraph("choudharyjairaj569@gmail.com | +91-8854098241")
doc.add_paragraph("LinkedIn: www.linkedin.com/in/jairaj-choudhary-156b29358 | GitHub: github.com/Jairaj238")

# Function: heading + line
def add_heading_with_line(text):
    doc.add_heading(text, level=1)
    p = doc.add_paragraph()
    run = p.add_run("―" * 60)   # ek line
    run.font.size = Pt(8)

# Career Objective
add_heading_with_line("Career Objective")
doc.add_paragraph(
    "I am a second-year student with a strong interest in learning new technologies and gaining hands-on experience. "
    "I am looking for opportunities where I can grow my skills, contribute to real-world projects, and become better "
    "prepared for my future career."
)

# Experience
add_heading_with_line("Experience")
doc.add_paragraph("Data Analytics Intern – Labmentix Pvt. Ltd. (Jun 2025 – Present)")
doc.add_paragraph("• Completed a 4-6 month internship focusing on data analytics and reporting.")
doc.add_paragraph("• Worked on real-world datasets to analyze trends, generate insights, and support decision-making.")
doc.add_paragraph("• Gained hands-on experience with data cleaning, visualization, and performance metrics.")
doc.add_paragraph("• Demonstrated consistency, problem-solving skills, and dedication throughout the internship.")

# Projects
add_heading_with_line("Projects")
doc.add_paragraph("Flights Data Analysis")
doc.add_paragraph("• Analyzed large U.S. flight datasets to study delays, cancellations, and airline performance using Pandas and SQL.")
doc.add_paragraph("• Developed KPIs and visualizations with Plotly to track on-time performance, delay reasons, and route efficiency.")
doc.add_paragraph("• Produced a cleaned dataset and summary report, enhancing data usability for deeper insights.")
doc.add_paragraph("Uber Demand & Supply Data Analysis")
doc.add_paragraph("• Conducted analysis of Uber ride data to identify demand–supply gaps across time and locations.")
doc.add_paragraph("• Derived insights on peak hours, cancellations, and driver availability using SQL and Excel.")
doc.add_paragraph("• Recommended strategies to optimize resource allocation and improve service efficiency.")
doc.add_paragraph("A Virtual Whiteboard")
doc.add_paragraph("• Developed a Virtual Whiteboard using HTML, CSS, and JavaScript.")
doc.add_paragraph("• Features include color selection, brush size control, and clear screen option.")
doc.add_paragraph("• Compatible with both mobile and desktop screens.")
doc.add_paragraph("• Learned how to handle mouse and touch events for drawing.")

# Education
add_heading_with_line("Education")
doc.add_paragraph("Indian Institute Of Technology (IIT) (2024 - Present)")
doc.add_paragraph("B.Sc (Hons) CSDA - CGPA: 8.74/10, Patna, Bihar")

# Certificates
add_heading_with_line("Certificates")
doc.add_paragraph("Python Programming Intern – Samyak IT Solutions Pvt. Ltd. (Mar 2025 – May 2025)")
doc.add_paragraph("Certification: Certified in Python Programming – NSDC")

# Technical Skills
add_heading_with_line("Technical Skills")
doc.add_paragraph("Programming Languages & Libraries: Python, SQL, NumPy, Pandas, Seaborn, Matplotlib")
doc.add_paragraph("Data Tools and Technologies: Power BI, Microsoft Excel, DBMS, Data Cleaning, Reporting & Dashboards, Statistical Analysis")
doc.add_paragraph("Core Analytical Skills: Data Analysis, Data Visualization, Data Interpretation, Problem Solving")
doc.add_paragraph("Web Designing: HTML, CSS, JavaScript")

# Save DOCX
docx_file = "Resume_with_lines.docx"
doc.save(docx_file)

# Convert to PDF
docx2pdf.convert(docx_file, "Resume_with_lines.pdf")

print("✅ Resume_with_lines.pdf generate ho gaya!")

python resume_with_lines.py
